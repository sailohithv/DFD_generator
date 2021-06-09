from docx import Document
from collections import OrderedDict
from docx.enum.text import WD_COLOR_INDEX
import DML_Parser_query_sep as dml_qs

sql_file = open("C:\\xampp\htdocs\SQL_parser\q_uploads\q_file.sql", "r")
file_to_string = sql_file.read()
temp=['tm.e_first_name', 'to.e_first_name', 'physician.e_first_name','cs.e_suppliername', 'cs.e_pin', 'nm.e_suppliername', 'nm.e_pin',
                          'suppliers.e_suppliername', 'suppliers.e_pin']
# print file_to_string

document = Document()

# paragraph = document.add_paragraph(file_to_string)
pattern_list=[]
position_list=[]
mis_itr_dict={}
for i in temp:
    x=file_to_string.lower().find(i)
    if(x>1):
        mis_itr_dict[x]=i
position_list=mis_itr_dict.keys()
position_list.sort()

temp_file_string=file_to_string
for it in position_list:
    # print t[it:it+len(dic[it])]
    x= file_to_string.split(temp_file_string[it:it + len(mis_itr_dict[it])])
    pattern_list.append(x[0])
    pattern_list.append(temp_file_string[it:it + len(mis_itr_dict[it])])
    file_to_string=x[-1]
pattern_list.append(temp_file_string[position_list[-1] + len(mis_itr_dict[position_list[-1]]):])
print pattern_list

for i in range(0, len(pattern_list)):
    pa = document.add_paragraph()
    if i%2==0:
        pa.add_run(pattern_list[i])
    else:
        font=pa.add_run(pattern_list[i]).font
        font.highlight_color = WD_COLOR_INDEX.RED

print document.save('fck.docx')



