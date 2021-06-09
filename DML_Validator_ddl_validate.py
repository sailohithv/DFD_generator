#!C:\Python27\python.exe
import cgi, os
import xlrd
from collections import defaultdict
from pandas import *
from collections import OrderedDict
from collections import Counter
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

print """\
Content-type:text/html\r\n\r\n

"""

import DML_Parser_main as dml_P
import DDL_Parser_ddl_main as ddl_P
import DML_Parser_query_sep as dml_qs

form = cgi.FieldStorage()
def dml_fetch():

    table_list=[]
    table_list.extend(column_names_dict.keys())
    # print table_list
    return table_list

def ddl_fecth(table_list):
    db_select=form.getvalue('db_select')
    out_list,db_type=ddl_P.DDL_parsing().get_user_selection(db_select,db_name,table_list)
    f_name=ddl_P.DDL_parsing().excel_write(out_list,db_type)
    return f_name

def excel_to_dict(file_path):
    xls = ExcelFile(file_path)
    df = xls.parse(xls.sheet_names[0])
    excel_dict = df.to_dict()
    table_name_dict = excel_dict['Table_name']
    col_name_dict = excel_dict['Column_name']
    tb_col_tuple = zip(table_name_dict.values(), col_name_dict.values())
    ddl_dict = defaultdict(list)
    for tuple_itr in tb_col_tuple:
        ddl_dict[tuple_itr[0].encode('utf-8')].append(tuple_itr[1].encode('utf-8'))
    return dict(ddl_dict)

def excel_read(path):
    loc = path

    wb = xlrd.open_workbook(loc)
    sheet = wb.sheet_by_index(0)
    sheet.cell_value(0, 0)

    tab_li = []
    di1 = {}
    for i in range(0, sheet.ncols):
        for cell in sheet.col(i):
            if i == 1:
                tab_li.append(cell.value.encode('utf-8'))
    tab_li = tab_li[1::]
    count_dict = Counter(tab_li)

    unique_tab_list = list(OrderedDict.fromkeys(tab_li))

    for key in count_dict:
        for i in unique_tab_list:
            if key == i:
                di1.setdefault(i, count_dict.get(key, ''))

    col_li = []
    for i in range(0, sheet.ncols):
        for cell in sheet.col(i):
            if i == 2:
                col_li.append(cell.value.encode('utf-8'))
    col_li = col_li[1::]

    dic = {}
    f = 0
    for i in di1:
        e = f + di1.get(i)
        dic.setdefault(i, col_li[f:e])
        f = e
    return dic

def compare_dict(dml_tbl_dict, ddl_tbl_dict):
    dif_dict = defaultdict(list)
    for tb_name, cl_names in dml_tbl_dict.items():
        # dif_dict[tb_name] = list(set(ddl_tbl_dict[tb_name]) - set(cl_names))
        for dml_col_itr in cl_names:
            # for ddl_col_itr in ddl_tbl_dict[tb_name]:
            if dml_col_itr not in ddl_tbl_dict[tb_name]:
                dif_dict[tb_name].append(dml_col_itr)
    return dict(dif_dict)


def mismatch_combination(dif_dict,alias_dict):
    mismatch_list=[]

    itr_dict=defaultdict(list)
    for tb_name in dif_dict.keys():
        for it in alias_dict[tb_name].values():
            if it =='--':
                continue
            else:
                itr_dict[tb_name].extend(it)
        itr_dict[tb_name].append(tb_name)
    for tb,tv in dif_dict.items():
        mismatch_list.extend([mis_alias + "." + mis_column for mis_alias in itr_dict[tb] for mis_column in tv])
    return mismatch_list

def word_download(mismatch_list):
    form = cgi.FieldStorage()
    document = Document()
    pattern_list = []
    position_list = []
    mis_itr_dict = {}
    sql_file = open(os.getcwd() + "\q_uploads\q_file.sql", "r")
    file_to_string = sql_file.read()
    temp_file_string = file_to_string

    for i in mismatch_list:
        x = file_to_string.lower().find(i)
        if (x > 1):
            mis_itr_dict[x] = i
    position_list = mis_itr_dict.keys()
    position_list.sort()

    for it in position_list:
        # print t[it:it+len(dic[it])]
        x= file_to_string.split(temp_file_string[it:it + len(mis_itr_dict[it])])
        pattern_list.append(x[0])
        pattern_list.append(temp_file_string[it:it + len(mis_itr_dict[it])])
        file_to_string=x[-1]
    pattern_list.append(temp_file_string[position_list[-1] + len(mis_itr_dict[position_list[-1]]):])

    for i in range(0, len(pattern_list)):
        pa = document.add_paragraph()
        if i%2==0:
            pa.add_run(pattern_list[i])
        else:
            font=pa.add_run(pattern_list[i]).font
            font.highlight_color = WD_COLOR_INDEX.RED
    filename='ErrorFile.docx'
    # print form.getvalue('filename')
    document.save("XlsFiles\\"+filename)
    return filename


if __name__== "__main__":
    db_name, tableName_dict, aliasName_dict, column_names_dict,table_with_query_dict = dml_P.dml_parser('q_file.sql')
    tab_list=dml_fetch()
    excel_file_name=ddl_fecth(tab_list)
    file_path=os.getcwd()+"\\XlsFiles\\"+excel_file_name
    ddl_tbl_dict=excel_to_dict(file_path)
    dif_dict=compare_dict(column_names_dict,ddl_tbl_dict)
    mismatch_list=mismatch_combination(dif_dict,aliasName_dict)
    filename=word_download(mismatch_list)
    print """
    <html>
            <head>
                  <script>
                  function download_file()
                  {
                     window.open('XlsFiles/"""+filename+"""')
                  }
                          </script>
                      </head>   
        <body align="center">
        <h3> DDL dictionary has been validated :"""+filename+"""</h3>
        <input type="submit" name="download_file" value="Download" onclick="download_file()" >
    """
    # print mismatch_list
    # print ddl_tbl_dict
    # print column_names_dict
    # print dif_dict
    print """</body>
    </html>
    """


