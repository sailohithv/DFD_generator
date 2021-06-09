from docx import Document



from docx import Document as doc
from docx.enum.text import WD_COLOR_INDEX
import DML_Parser_query_sep as dml_qs
from collections import OrderedDict

#
sql_file = open("C:\\xampp\htdocs\SQL_parser\q_uploads\q_file.sql", "r")
file_to_string = sql_file.read()

# file_to_string=re.sub("\s\s*", " ", file_to_string).lower()
queries = file_to_string.lower().split(';')
tb,tb_with_query= dml_qs.query_processing().table_nameList_query(queries)

#
# col_list = {'physician': ['tm.e_first_name', 'to.e_first_name', 'physician.e_first_name'],
#             'suppliers': ['cs.e_suppliername', 'cs.e_pin', 'nm.e_suppliername', 'nm.e_pin',
#                           'suppliers.e_suppliername', 'suppliers.e_pin']}
col_list = {'physician': ['to.e_first_name'],'suppliers': ['cs.e_suppliername', 'cs.e_pin']}
temp=['tm.e_first_name', 'to.e_first_name', 'physician.e_first_name','cs.e_suppliername', 'cs.e_pin', 'nm.e_suppliername', 'nm.e_pin',
                          'suppliers.e_suppliername', 'suppliers.e_pin']

dc = doc('C:\\xampp\htdocs\SQL_parser\\final_test.docx')
print len(dc.paragraphs)
print dc.paragraphs[0].text


doc1 = Document()



for tb_name,query in tb_with_query.items():
    p1 = doc1.add_paragraph()
    if tb_name in col_list.keys():
        for m_itr in temp:
            query_str=query.split(m_itr)
            for sub in query_str[:-1]:
                p1.add_run(sub)
                font=p1.add_run(m_itr).font
                font.highlight_color = WD_COLOR_INDEX.RED
            p1.add_run(query_str[-1])
    else:
        p1.add_run(query)
doc1.save('t.docx')
